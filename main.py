from docx import Document

# Create a new Word document
doc = Document()

# -------------------------------
# Title
# -------------------------------
doc.add_heading('Lab Activity-03: CRISP-DM Reflections', 0)

# -------------------------------
# Step-2 Reflection
# -------------------------------
doc.add_heading('Step-2: Data Understanding', level=1)

# Q1
doc.add_heading('1. Challenges during MySQL Data-dump Restoration (5 pts)', level=2)
doc.add_paragraph(
    "The main challenge I faced during the restoration was that the 'mysql' command "
    "was not recognized in cmd.exe. I solved this by navigating to the MySQL 'bin' "
    "directory (C:\\Program Files\\MySQL\\MySQL Server 8.0\\bin) and running the command "
    "from there. After that, I successfully imported the dump file into the database."
)

# Q2
doc.add_heading('2. Table Details (5 pts)', level=2)
doc.add_paragraph("Number of Attributes: __________________")
doc.add_paragraph("Number of Objects: __________________")

# Q3
doc.add_heading('3. Attribute Data Types (10 pts)', level=2)
table = doc.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Attribute'
hdr_cells[1].text = 'Data Type'

# Add placeholder rows
attributes = [
    ('id', 'Integer'),
    ('age', 'Integer'),
    ('job', 'Varchar'),
    ('marital', 'Varchar'),
    ('education', 'Varchar'),
    ('...', '...')
]

for attr, dtype in attributes:
    row_cells = table.add_row().cells
    row_cells[0].text = attr
    row_cells[1].text = dtype

# -------------------------------
# Step-3 Reflection
# -------------------------------
doc.add_heading('Step-3: Data Preparation', level=1)

# Q1 Initial summary
doc.add_heading('1. Initial Summary (10 pts)', level=2)
doc.add_paragraph("a. Job\nb. Marital\nc. Education\nd. Contact\n\n[Paste SQL frequency results here]")

# Q2 Final summary
doc.add_heading('2. Final Summary (15 pts)', level=2)
doc.add_paragraph("a. Job\nb. Marital\nc. Education\nd. Contact\n\n[Paste SQL frequency results after cleaning here]")

# Q3 Issues during clean-up
doc.add_heading('3. Issues during Clean-up (5 pts)', level=2)
doc.add_paragraph(
    "Some issues encountered include inconsistent capitalization and spelling, multiple "
    "variants of the same category (e.g., 'admin', 'administrator'), missing values labeled "
    "as 'unknown' or 'na', and deciding a threshold for grouping rare categories."
)

# Q4 Repercussions if no clean-up
doc.add_heading('4. Repercussions if Clean-up is Neglected (10 pts)', level=2)
doc.add_paragraph(
    "If clean-up is neglected, the dataset will contain noise and redundant categories "
    "which can reduce model accuracy. It may also cause inflated feature dimensionality, "
    "lead to poor generalization, and produce misleading business insights."
)

# -------------------------------
# Save to your Week8Lab folder
# -------------------------------
output_path = r"C:\Users\Lyne\Desktop\Week8Lab\Lab_Activity_Reflection_Template.docx"
doc.save(output_path)

print(f"Template saved successfully at: {output_path}")
